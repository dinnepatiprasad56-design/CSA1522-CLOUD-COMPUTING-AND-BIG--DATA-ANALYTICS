"""
Convert Individual_Contribution_Report.md to a beautifully styled,
professional academic DOCX document.
"""

import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) for a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    """Apply clean borders to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def make_row_cant_split(row):
    """Ensure table row does not split across pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def make_row_header(row):
    """Mark row as repeat table header across pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def add_footer_page_number(run):
    """Insert a dynamic page number field in the footer."""
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_footer_total_pages(run):
    """Insert total page count field in the footer."""
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def clean_math_markdown(text):
    """Convert LaTeX math fragments into clean representations."""
    text = text.replace(r"$\ge$", "≥").replace(r"$\ge 42.0$", "≥ 42.0").replace(r"$\ge 65.0$", "≥ 65.0")
    text = text.replace(r"$\ge 60.0$", "≥ 60.0").replace(r"$\le 5.0$", "≤ 5.0").replace(r"$\le 15.0$", "≤ 15.0")
    text = text.replace(r"$\le 970.0$", "≤ 970.0").replace(r"$\le$", "≤").replace(r"$\ge$", "≥")
    text = text.replace(r"$O(N)$", "O(N)").replace(r"$O(1)$", "O(1)")
    text = text.replace(r"$\le 50\,\text{MB}$", "≤ 50 MB")
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    return text

def parse_inline_formatting(paragraph, text, base_font_size=11, base_color=(0, 0, 0), italic_default=False):
    """
    Parse inline bold, italic, code, placeholders, and markdown links
    into runs on the given paragraph.
    """
    text = clean_math_markdown(text)
    
    token_pattern = re.compile(
        r'(\*\*\*[^*]+\*\*\*'
        r'|\*\*[^*]+\*\*'
        r'|\*[^*]+\*'
        r'|`[^`]+`'
        r'|\[INSERT [^\]]+\]'
        r'|\[[^\]]+\]\([^)]+\))'
    )
    
    pos = 0
    for match in token_pattern.finditer(text):
        start, end = match.span()
        if start > pos:
            run = paragraph.add_run(text[pos:start])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(base_font_size)
            run.font.color.rgb = RGBColor(*base_color)
            run.font.italic = italic_default
        
        token = match.group(0)
        
        if token.startswith('***') and token.endswith('***'):
            run = paragraph.add_run(token[3:-3])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(base_font_size)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = RGBColor(*base_color)
        elif token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(base_font_size)
            run.font.bold = True
            run.font.italic = italic_default
            run.font.color.rgb = RGBColor(*base_color)
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(base_font_size)
            run.font.italic = True
            run.font.color.rgb = RGBColor(*base_color)
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(base_font_size - 1)
            run.font.color.rgb = RGBColor(180, 40, 40)
            run.font.bold = False
        elif token.startswith('[INSERT ') and token.endswith(']'):
            run = paragraph.add_run(token)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(base_font_size)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = RGBColor(192, 0, 0) # Dark red for action items
        elif token.startswith('[') and '](' in token and token.endswith(')'):
            link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', token)
            if link_match:
                link_text, link_url = link_match.groups()
                run = paragraph.add_run(link_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(base_font_size)
                run.font.color.rgb = RGBColor(31, 78, 121)
                run.font.underline = True
            else:
                run = paragraph.add_run(token)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(base_font_size)
                run.font.color.rgb = RGBColor(*base_color)
        
        pos = end
    
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(base_font_size)
        run.font.color.rgb = RGBColor(*base_color)
        run.font.italic = italic_default


def render_code_block(doc, code_lines, lang=""):
    """Render a clean, shaded code block in a single-cell container."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.27)
    set_cell_background(cell, "F8F9FA")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="1F4E79"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="D0D7DE"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    
    full_text = "\n".join(code_lines)
    run = p.add_run(full_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.0)
    run.font.color.rgb = RGBColor(36, 41, 47)
    
    sp_p = doc.add_paragraph()
    sp_p.paragraph_format.space_before = Pt(0)
    sp_p.paragraph_format.space_after = Pt(4)


def calculate_proportional_widths(normalized_rows, total_width_inches=6.27):
    """Compute balanced column widths based on maximum text length in each column."""
    col_count = len(normalized_rows[0])
    max_lens = [0] * col_count
    
    for row in normalized_rows:
        for c_idx, cell_text in enumerate(row):
            max_lens[c_idx] = max(max_lens[c_idx], len(str(cell_text).strip()))
    
    # Enforce minimum weight per column
    weights = [max(l, 8) for l in max_lens]
    # If a column is overwhelmingly long, soften its weight power to avoid starving other columns
    softened_weights = [w ** 0.85 for w in weights]
    total_weight = sum(softened_weights)
    
    widths = [(w / total_weight) * total_width_inches for w in softened_weights]
    # Ensure no column is less than 0.6 inches
    adjusted_widths = [max(w, 0.6) for w in widths]
    scale = total_width_inches / sum(adjusted_widths)
    return [Inches(w * scale) for w in adjusted_widths]


def render_markdown_table(doc, table_rows):
    """Parse and render a clean, professional academic table in Word."""
    if not table_rows:
        return
    
    parsed_rows = []
    for r in table_rows:
        cells = [c.strip() for c in r.strip().strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells if c):
            continue
        parsed_rows.append(cells)
    
    if not parsed_rows:
        return
    
    col_count = max(len(row) for row in parsed_rows)
    normalized_rows = []
    for row in parsed_rows:
        if len(row) < col_count:
            row = row + [""] * (col_count - len(row))
        normalized_rows.append(row[:col_count])
    
    table = doc.add_table(rows=len(normalized_rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="D3D3D3", sz="4")
    
    col_widths = calculate_proportional_widths(normalized_rows, total_width_inches=6.27)
    
    for row_idx, row_data in enumerate(normalized_rows):
        row = table.rows[row_idx]
        make_row_cant_split(row)
        
        is_header = (row_idx == 0)
        if is_header:
            make_row_header(row)
            bg_color = "1F4E79" # Professional Navy Blue Header
            text_color = (255, 255, 255) # White text
        else:
            bg_color = "F9FAFB" if (row_idx % 2 == 1) else "FFFFFF"
            text_color = (0, 0, 0)
        
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=130, right=130)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            
            if is_header:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                parse_inline_formatting(p, f"**{cell_text}**", base_font_size=9.5, base_color=text_color)
            else:
                # Center short codes, numbers, statuses; left-align descriptive text
                if len(cell_text) <= 12 or cell_text in ["PASS", "FAIL", "Achieved", "Partial", "Completed"]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                parse_inline_formatting(p, cell_text, base_font_size=9.0, base_color=text_color)
    
    sp_p = doc.add_paragraph()
    sp_p.paragraph_format.space_before = Pt(0)
    sp_p.paragraph_format.space_after = Pt(6)


def convert_report_md_to_docx(md_path, docx_path):
    print(f"Reading markdown report from: {md_path}")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    
    doc = docx.Document()
    
    # Page Setup - A4 Standard for Engineering College Reports
    for section in doc.sections:
        section.page_width = Inches(8.27)   # 210 mm
        section.page_height = Inches(11.69) # 297 mm
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Unlink header/footer on title page
        section.different_first_page_header_footer = True
        
        # Header setup (pages 2+)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Individual Contribution Report | CSA1522 Cloud Computing & Big Data Analytics")
        hrun.font.name = 'Times New Roman'
        hrun.font.size = Pt(8.5)
        hrun.font.italic = True
        hrun.font.color.rgb = RGBColor(128, 128, 128)
        
        # Footer setup (pages 2+)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        frun = fp.add_run("SIMATS Engineering | Department of CSE    —    Page ")
        frun.font.name = 'Times New Roman'
        frun.font.size = Pt(9.0)
        frun.font.color.rgb = RGBColor(100, 100, 100)
        add_footer_page_number(frun)
        frun2 = fp.add_run(" of ")
        frun2.font.name = 'Times New Roman'
        frun2.font.size = Pt(9.0)
        frun2.font.color.rgb = RGBColor(100, 100, 100)
        add_footer_total_pages(frun2)
    
    # Configure Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(20, 20, 20)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)
    
    in_code_block = False
    code_lines = []
    code_lang = ""
    
    in_table = False
    table_lines = []
    
    i = 0
    total_lines = len(lines)
    
    while i < total_lines:
        line = lines[i].rstrip('\r\n')
        stripped = line.strip()
        
        # Code Block Handler
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                render_code_block(doc, code_lines, code_lang)
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Table Handler
        if stripped.startswith("|") and stripped.endswith("|"):
            if not in_table:
                in_table = True
                table_lines = [stripped]
            else:
                table_lines.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                render_markdown_table(doc, table_lines)
                table_lines = []
        
        # Page Break Handler
        if stripped == "\\newpage":
            doc.add_page_break()
            i += 1
            continue
        
        # Empty Line
        if not stripped:
            i += 1
            continue
        
        # Horizontal Rule Handler (---)
        if stripped in ["---", "___", "***"]:
            i += 1
            continue
        
        # Headings
        if stripped.startswith("# "):
            h_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            if "INDIVIDUAL CONTRIBUTION" in h_text or "CERTIFICATE" in h_text or "DECLARATION" in h_text:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            parse_inline_formatting(p, f"**{h_text}**", base_font_size=17, base_color=(0, 32, 96))
            i += 1
            continue
        
        if stripped.startswith("## "):
            h_text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(13)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            if "HADOOP STREAMING-BASED" in h_text:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            parse_inline_formatting(p, f"**{h_text}**", base_font_size=13.5, base_color=(31, 78, 121))
            i += 1
            continue
        
        if stripped.startswith("### "):
            h_text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            
            if "A Capstone Project Report" in h_text:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            parse_inline_formatting(p, f"**{h_text}**", base_font_size=12, base_color=(46, 117, 182))
            i += 1
            continue
        
        if stripped.startswith("#### "):
            h_text = stripped[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            parse_inline_formatting(p, f"***{h_text}***", base_font_size=11, base_color=(46, 117, 182))
            i += 1
            continue
        
        # Blockquote Handler
        if stripped.startswith("> "):
            bq_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.35)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            parse_inline_formatting(p, bq_text, base_font_size=10.5, base_color=(60, 60, 60), italic_default=True)
            i += 1
            continue
        
        # Bullet Lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            item_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.12
            
            run_bullet = p.add_run("•  ")
            run_bullet.font.name = 'Times New Roman'
            run_bullet.font.size = Pt(10)
            run_bullet.font.color.rgb = RGBColor(31, 78, 121)
            run_bullet.font.bold = True
            
            parse_inline_formatting(p, item_text, base_font_size=11, base_color=(20, 20, 20))
            i += 1
            continue
        
        # Numbered Lists
        num_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if num_match:
            num_val, item_text = num_match.groups()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.12
            
            run_num = p.add_run(f"{num_val}.  ")
            run_num.font.name = 'Times New Roman'
            run_num.font.size = Pt(10.5)
            run_num.font.bold = True
            run_num.font.color.rgb = RGBColor(31, 78, 121)
            
            parse_inline_formatting(p, item_text, base_font_size=11, base_color=(20, 20, 20))
            i += 1
            continue
        
        # Special Centered Lines for Title Page
        title_center_phrases = [
            "Submitted in partial fulfilment",
            "Bachelor of Engineering",
            "Computer Science and Engineering",
            "Department of Computer Science and Engineering",
            "SIMATS Engineering",
            "Dinnepati Sindhu Prasad",
            "192311271",
            "Place: Chennai",
            "Date: [DATE]"
        ]
        
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        if any(phrase in stripped for phrase in title_center_phrases) and i < 110:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        parse_inline_formatting(p, stripped, base_font_size=11, base_color=(20, 20, 20))
        i += 1
    
    if in_table:
        render_markdown_table(doc, table_lines)
    
    print(f"Saving formatted Word document to: {docx_path}")
    doc.save(docx_path)
    print("Successfully created enhanced DOCX report!")

if __name__ == "__main__":
    base_dir = r"c:\Users\prasa\Downloads\CSA1522-CLOUD-COMPUTING-AND-BIG--DATA-ANALYTICS\Capstone Project"
    md_file = os.path.join(base_dir, "Individual_Contribution_Report.md")
    docx_file = os.path.join(base_dir, "Individual_Contribution_Report.docx")
    convert_report_md_to_docx(md_file, docx_file)
